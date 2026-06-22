from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
README = ROOT / "README_zh.md"
OUT_DIR = ROOT


def read_source() -> str:
    return README.read_text(encoding="utf-8")


def section(markdown: str, heading: str, next_headings: list[str] | None = None) -> str:
    start = markdown.find(heading)
    if start < 0:
        return ""
    end = len(markdown)
    if next_headings:
        positions = [markdown.find(h, start + len(heading)) for h in next_headings]
        positions = [p for p in positions if p >= 0]
        if positions:
            end = min(positions)
    return markdown[start:end].strip()


def strip_html_blocks(text: str) -> str:
    text = re.sub(r"<div[^>]*>|</div>", "", text)
    text = re.sub(r"<h1>|</h1>", "", text)
    text = re.sub(r'<img\s+src="([^"]+)"\s+alt="([^"]*)"\s*/?>', r"图片：\2（\1）", text)
    return text


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.strip()


def set_east_asia_font(run, font_name="SimSun"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(clean_inline(text))
    set_east_asia_font(r)


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.style = "Code"
    for idx, line in enumerate(code.rstrip("\n").splitlines()):
        if idx:
            p.add_run().add_break(WD_BREAK.LINE)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.size = Pt(9)


def maybe_add_table(doc: Document, lines: list[str]) -> bool:
    if len(lines) < 2 or not all(line.strip().startswith("|") for line in lines):
        return False
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return False
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            set_cell_text(table.rows[r_idx].cells[c_idx], row[c_idx] if c_idx < len(row) else "")
    return True


def add_markdown(doc: Document, markdown: str):
    markdown = strip_html_blocks(markdown)
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            if not maybe_add_table(doc, table_lines):
                for line in table_lines:
                    doc.add_paragraph(clean_inline(line))
            table_lines = []

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            text = clean_inline(stripped.lstrip("#").strip())
            doc.add_heading(text, level=level)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        else:
            doc.add_paragraph(clean_inline(stripped))
    flush_table()
    if code_lines:
        add_code_block(doc, "\n".join(code_lines))


def configure_doc(doc: Document):
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10.5)
    if "Code" not in styles:
        styles.add_style("Code", 1)
    styles["Code"].font.name = "Consolas"
    styles["Code"].font.size = Pt(9)


def write_docx(filename: str, title: str, intro: str, body: str):
    doc = Document()
    configure_doc(doc)
    doc.add_heading(title, level=0)
    doc.add_paragraph(intro)
    add_markdown(doc, body)
    doc.core_properties.title = title
    doc.core_properties.subject = "DISCOVERSE README_zh.md 拆分文档"
    doc.core_properties.author = "Codex"
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_documents():
    md = read_source()

    install = section(md, "## 📦 安装与快速开始", ["## 📷 高保真渲染设置"])
    rendering = section(md, "## 📷 高保真渲染设置", ["## 🔨 Real2Sim管道"])
    real2sim = section(md, "## 🔨 Real2Sim管道", ["## 💡 使用示例"])
    examples = section(md, "## 💡 使用示例", ["## 🎓 学习与训练"])
    training = section(md, "## 🎓 学习与训练", ["## ⏩ 最近更新"])
    updates = section(md, "## ⏩ 最近更新", ["## ❔ 故障排除"])
    troubleshooting = section(md, "## ❔ 故障排除", ["## ⚖️ 许可证"])

    dev_body = "\n\n".join(
        [
            "## 开发环境与工具版本概览\n"
            "- 操作系统：README 提供 Linux 与 macOS 的 Git LFS 安装方式。\n"
            "- Python：推荐使用 conda 创建 discoverse 环境，Python 版本为 3.10，README 标注 Python >= 3.8 即可。\n"
            "- 包管理与源码安装：使用 pip install -e . 进行可编辑安装。\n"
            "- Git LFS：用于拉取大文件资源。\n"
            "- CUDA：高保真 3DGS 渲染建议安装 CUDA 11.8+，并根据显卡驱动选择匹配版本。\n"
            "- Docker/GPU 容器：Docker 部署需要 NVIDIA Container Toolkit。\n",
            install,
            rendering,
        ]
    )

    deploy_body = "\n\n".join(
        [
            "## 部署范围\n"
            "本部分覆盖源码部署、Docker 部署、模型资源准备、基础运行验证和常用示例命令。",
            install,
            rendering,
            real2sim,
            examples,
            training,
        ]
    )

    ops_body = "\n\n".join(
        [
            "## 运维计划\n"
            "- 环境基线：固定 Python、CUDA、Docker 镜像 tag 与关键可选依赖组合，避免开发机和部署机依赖漂移。\n"
            "- 安装验证：每次新环境部署后运行 python scripts/check_installation.py，并至少执行一个基础仿真示例。\n"
            "- 镜像维护：Docker 镜像建议按版本 tag 管理；README 当前提到预构建镜像更新至 v1.8.6。\n"
            "- 模型资源：3DGS 模型默认位于 models/3dgs，首次运行会自动下载；国内网络环境可配置 HF_ENDPOINT 使用镜像。\n"
            "- GPU 与显示：容器部署时检查 NVIDIA Container Toolkit、--gpus all、DISPLAY、X11 挂载和 xhost 权限。\n"
            "- 故障闭环：安装或运行异常优先查阅 discoverse/doc/troubleshooting.md，并记录问题现象、环境版本、复现命令和修复方式。\n",
            "## 使用建议\n"
            "- 初学或基础开发：先安装基础模块 pip install -e .，确认核心仿真可运行后再增加可选模块。\n"
            "- LiDAR/SLAM：选择 lidar 与 visualization 额外依赖，关注 Taichi 与 GPU 环境兼容性。\n"
            "- 模仿学习：选择 act_full 或对应算法依赖，按数据收集、训练、推理三个阶段管理数据和模型产物。\n"
            "- 高保真视觉仿真：选择 gs 依赖并准备 CUDA 11.8+；无 3DGS 需求或已使用 Docker 时可跳过本机渲染配置。\n"
            "- 远程图形环境：优先考虑 Dockerfile.vnc 版本，减少本地显示服务器依赖。\n",
            updates,
            troubleshooting,
        ]
    )

    return [
        write_docx(
            "DISCOVERSE_开发环境与工具版本及其安装说明.docx",
            "DISCOVERSE 开发环境与工具版本及其安装说明",
            "本文档根据 README_zh.md 中的安装、依赖、工具版本和渲染环境内容整理生成。",
            dev_body,
        ),
        write_docx(
            "DISCOVERSE_部署安装与使用说明.docx",
            "DISCOVERSE 部署安装与使用说明",
            "本文档根据 README_zh.md 中的部署安装、Docker、模型准备、使用示例和训练流程内容整理生成。",
            deploy_body,
        ),
        write_docx(
            "DISCOVERSE_运维计划和使用建议.docx",
            "DISCOVERSE 运维计划和使用建议",
            "本文档基于 README_zh.md 的更新、安装验证、故障排除和典型使用场景整理生成。",
            ops_body,
        ),
    ]


if __name__ == "__main__":
    for p in build_documents():
        print(p)
